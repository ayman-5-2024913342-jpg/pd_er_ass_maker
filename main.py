import os
import subprocess
import logging
import re
import time
import shutil
from git import Repo
from docx import Document
from docx.shared import Pt, RGBColor
from pygments import lexers, lex
from groq import Groq, RateLimitError

# --- Configuration ---
GROQ_API_KEY = "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"
REPO_URL = "https://github.com/ayman-5-2024913342-jpg/Ass.git"
MODEL_NAME = "llama-3.1-8b-instant"
WORKING_DIR = "./fortran_workspace"
OUTPUT_DOC = "Assignments_Final_Report.docx"

client = Groq(api_key=GROQ_API_KEY)
logging.basicConfig(level=logging.INFO, format='%(asctime)s [%(levelname)s] %(message)s')
logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Groq helpers
# ---------------------------------------------------------------------------

def groq_call(prompt: str, system: str = "You are a helpful assistant.") -> str:
    """Call Groq with automatic rate-limit retry."""
    time.sleep(1.5)
    try:
        completion = client.chat.completions.create(
            model=MODEL_NAME,
            messages=[
                {"role": "system", "content": system},
                {"role": "user",   "content": prompt},
            ],
            temperature=0.0,
        )
        return completion.choices[0].message.content.strip()
    except RateLimitError:
        logger.warning("Rate limit hit – sleeping 30 s …")
        time.sleep(30)
        return groq_call(prompt, system)


# ---------------------------------------------------------------------------
# Input / file detection
# ---------------------------------------------------------------------------

def detect_input_files(code: str) -> list[str]:
    """
    Return a list of filenames that the Fortran source tries to OPEN for reading.
    We look for  OPEN(... FILE='name' ...)  patterns (case-insensitive).
    """
    pattern = re.compile(
        r"open\s*\([^)]*file\s*=\s*['\"]([^'\"]+)['\"]",
        re.IGNORECASE,
    )
    return list(dict.fromkeys(pattern.findall(code)))  # deduplicated, order-preserved


def detect_stdin_reads(code: str) -> bool:
    """Return True if the code reads from standard input (READ(*,*) etc.)."""
    return bool(re.search(r"\bread\s*\(\s*\*\s*,", code, re.IGNORECASE))


def generate_input_file(code: str, filename: str, folder_path: str) -> None:
    """
    Ask Groq to produce realistic numeric content for *filename* as required
    by the Fortran source, then write it to disk.
    """
    prompt = (
        f"The following Fortran program reads data from a file named '{filename}'.\n"
        f"Generate ONLY the raw file content (numbers/strings, one per line as the "
        f"program expects). No explanation, no code fences.\n\n"
        f"CODE:\n{code}"
    )
    content = groq_call(prompt, system="You are a Fortran I/O specialist. Return ONLY raw file data.")
    # Strip any accidental markdown fences
    content = re.sub(r"```[^\n]*\n?", "", content).strip()
    dest = os.path.join(folder_path, filename)
    with open(dest, "w") as fh:
        fh.write(content + "\n")
    logger.info(f"    Created input file '{filename}' ({len(content.splitlines())} lines)")


def generate_stdin_input(code: str) -> str:
    """
    Ask Groq for the exact stdin values needed to run the program.
    Returns a newline-separated string of values, or '' if none needed.
    """
    prompt = (
        "The following Fortran program reads from standard input (READ(*,*)).\n"
        "Return ONLY the input values the program needs, one per line, in the exact "
        "order they are read. Use simple numeric/string literals. "
        "No explanation, no labels, no code fences.\n"
        "If no stdin input is needed, return exactly: NONE\n\n"
        f"CODE:\n{code}"
    )
    result = groq_call(prompt, system="You are a Fortran I/O specialist. Return ONLY raw input data.")
    result = re.sub(r"```[^\n]*\n?", "", result).strip()
    if result.upper() == "NONE" or not result:
        return ""
    return result


# ---------------------------------------------------------------------------
# Compilation & execution
# ---------------------------------------------------------------------------

def compile_fortran(source_path: str, exe_path: str) -> tuple[bool, str]:
    """Compile with gfortran. Returns (success, stderr)."""
    proc = subprocess.run(
        ["gfortran", source_path, "-o", exe_path],
        capture_output=True, text=True,
    )
    return proc.returncode == 0, proc.stderr


def fix_and_recompile(code: str, error: str, source_path: str, exe_path: str) -> str:
    """Ask Groq to fix syntax errors then recompile. Returns (possibly fixed) code."""
    logger.warning("  Attempting AI-assisted fix …")
    fix_prompt = (
        f"Fix ALL syntax errors in this Fortran code. Return ONLY the corrected code, "
        f"no explanation, no markdown fences.\n\nCODE:\n{code}\n\nERROR:\n{error}"
    )
    fixed = groq_call(fix_prompt)
    fixed = re.sub(r"```[^\n]*\n?", "", fixed).strip()
    with open(source_path, "w") as fh:
        fh.write(fixed)
    ok, _ = compile_fortran(source_path, exe_path)
    return fixed if ok else code


def run_program(executable: str, stdin_data: str, folder_path: str) -> str:
    """Execute a compiled Fortran binary and return its output (stdout or stderr)."""
    try:
        proc = subprocess.run(
            [executable],
            input=stdin_data if stdin_data else None,
            capture_output=True,
            text=True,
            timeout=10,
            cwd=folder_path,
        )
        output = proc.stdout.strip()
        if not output:
            output = proc.stderr.strip()
        return output or "(no output)"
    except subprocess.TimeoutExpired:
        return "ERROR: Program timed out (>10 s). It may be waiting for more input."
    except Exception as exc:
        return f"ERROR: {exc}"


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def add_highlighted_code(doc: Document, code: str) -> Document:
    """Add syntax-highlighted Fortran code inside a single-cell table."""
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    paragraph = cell.paragraphs[0]

    lexer = lexers.get_lexer_by_name("fortran")
    token_colors = {
        "Keyword": RGBColor(0, 0, 200),
        "Comment": RGBColor(0, 128, 0),
        "String":  RGBColor(163, 21, 21),
    }
    for ttype, value in lex(code, lexer):
        run = paragraph.add_run(value)
        run.font.name = "Consolas"
        run.font.size = Pt(9)
        for key, rgb in token_colors.items():
            if key in str(ttype):
                run.font.color.rgb = rgb
    return doc


def add_output_block(doc: Document, label: str, text: str) -> None:
    """Add a monospaced captioned output block."""
    doc.add_paragraph(label, style="Caption")
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


# ---------------------------------------------------------------------------
# Main pipeline
# ---------------------------------------------------------------------------

def process_file(f90_path: str, folder_path: str, doc: Document) -> None:
    name = os.path.splitext(os.path.basename(f90_path))[0]
    logger.info(f"  Processing {name}.f90 …")

    with open(f90_path, "r", errors="replace") as fh:
        code = fh.read()

    exe_path = os.path.join(folder_path, f"{name}.out")

    # --- Step 1: Compile ---
    ok, err = compile_fortran(f90_path, exe_path)
    if not ok:
        logger.warning(f"  Compilation failed for {name}, trying AI fix …")
        code = fix_and_recompile(code, err, f90_path, exe_path)
        ok, err = compile_fortran(f90_path, exe_path)

    # --- Step 2: Detect & create input files ---
    input_files = detect_input_files(code)
    for fname in input_files:
        fpath = os.path.join(folder_path, fname)
        if not os.path.exists(fpath):
            generate_input_file(code, fname, folder_path)

    # --- Step 3: Detect & generate stdin ---
    stdin_data = ""
    if detect_stdin_reads(code):
        stdin_data = generate_stdin_input(code)
        if stdin_data:
            logger.info(f"  Stdin values: {repr(stdin_data[:80])}")
        else:
            logger.info("  No stdin values needed (READ(*,*) present but Groq returned NONE)")

    # --- Step 4: Run ---
    output = "(compilation failed – skipped execution)"
    if ok and os.path.exists(exe_path):
        output = run_program(exe_path, stdin_data, folder_path)
        logger.info(f"  Output: {output[:120]}")

        # If output looks like a hang/timeout and we have stdin, warn
        if "waiting" in output.lower() or "timeout" in output.lower():
            logger.warning(f"  Program may need more stdin values for {name}")

    # --- Step 5: Write to document ---
    doc.add_heading(name, level=2)
    add_highlighted_code(doc, code)

    meta_parts = []
    if stdin_data:
        meta_parts.append(f"STDIN: {stdin_data.replace(chr(10), ' | ')}")
    if input_files:
        meta_parts.append(f"Input files: {', '.join(input_files)}")
    if not meta_parts:
        meta_parts.append("No external input required")
    doc.add_paragraph(" | ".join(meta_parts), style="Caption")

    add_output_block(doc, "Program Output:", output)

    # Cleanup
    if os.path.exists(exe_path):
        os.remove(exe_path)


def main() -> None:
    # Clean workspace
    if os.path.exists(WORKING_DIR):
        logger.info("Cleaning existing workspace …")
        shutil.rmtree(WORKING_DIR, ignore_errors=True)

    logger.info(f"Cloning {REPO_URL} …")
    try:
        Repo.clone_from(
            REPO_URL,
            WORKING_DIR,
            multi_options=["-c core.protectNTFS=false"],
            allow_unsafe_options=True,
        )
    except Exception as exc:
        logger.error(f"Clone failed: {exc}")
        return

    doc = Document()
    doc.add_heading("Fortran Automation Report", 0)

    assignment_folders = sorted(
        d for d in os.listdir(WORKING_DIR)
        if os.path.isdir(os.path.join(WORKING_DIR, d)) and d.lower().startswith("assignment")
    )

    if not assignment_folders:
        logger.error("No Assignment* folders found. Check repo structure.")
        return

    for folder in assignment_folders:
        folder_path = os.path.abspath(os.path.join(WORKING_DIR, folder))
        logger.info(f"=== {folder} ===")
        doc.add_heading(folder, level=1)

        f90_files = sorted(
            f for f in os.listdir(folder_path) if f.endswith(".f90")
        )
        if not f90_files:
            doc.add_paragraph("(no .f90 files found)")
            continue

        for f90 in f90_files:
            process_file(os.path.join(folder_path, f90), folder_path, doc)

    doc.save(OUTPUT_DOC)
    logger.info(f"Report saved → {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
